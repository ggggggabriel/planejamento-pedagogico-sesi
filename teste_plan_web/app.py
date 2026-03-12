from flask import Flask, render_template, request, send_file
from docx import Document
from werkzeug.utils import secure_filename
import pandas as pd
import io
import os
import sys

app = Flask(__name__)

EXCEL_FILE = 'HABILIDADES FUND.xlsx'

try:
    df = pd.read_excel(EXCEL_FILE)
    disciplinas = df.columns.tolist()
except FileNotFoundError:
    print(f"Erro: arquivo '{EXCEL_FILE}' não encontrado.")
    sys.exit(1)

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/create-plan', methods=['GET', 'POST'])
def create_plan():
    if request.method == 'POST':
        unidade = request.form.get('unidade', '')
        professor = request.form.get('professor', '')
        serie = request.form.get('serie', '')
        bimestre = request.form.get('bimestre', '')
        periodo = request.form.get('periodo', '')
        a = request.form.get('a', '')
        capitulo = request.form.get('capitulo', '')
        disciplina = request.form.get('disciplina', '')
        habilidades = request.form.get('habilidades', '')
        desenvolvimento = request.form.get('desenvolvimento', '')
        estrategia = request.form.get('estrategia', '')
        observacoes = request.form.get('observacoes', '')
        duracao_aula = request.form.get('duracao_aula', '')
        nome_arquivo = secure_filename(request.form.get('nome_arquivo', 'plano')) or 'plano'

        buffer = inserir_texto_no_modelo({
            'Unidade': unidade,
            'Professor': professor,
            'Série': serie,
            'Bimestre': bimestre,
            'Período': periodo,
            'Período Fim': a,
            'Capítulo': capitulo,
            'Área ou Disciplina': disciplina,
            'Habilidades': habilidades,
            'Desenvolvimento da Aula': desenvolvimento,
            'Estratégia': estrategia,
            'Observações': observacoes,
            'Duração da Aula': duracao_aula
        }, 'static/docx/Folha do Planejamento Pedagógico.docx')

        if buffer:
            return send_file(
                buffer,
                as_attachment=True,
                download_name=f"{nome_arquivo}.docx",
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        else:
            return "Erro ao gerar o arquivo", 500

    return render_template('form.html', disciplinas=disciplinas)

@app.route('/get-habilidades')
def get_habilidades():
    disciplina = request.args.get('disciplina')
    if disciplina in df.columns:
        habilidades = df[disciplina].dropna().tolist()
        return {'habilidades': habilidades}
    return {'habilidades': []}

def inserir_texto_no_modelo(dados, modelo_path):
    try:
        doc = Document(modelo_path)

        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        for chave, valor in dados.items():
            row = table.add_row()
            label_run = row.cells[0].paragraphs[0].add_run(chave)
            label_run.bold = True
            row.cells[1].paragraphs[0].add_run(str(valor))

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        print(f"Erro ao inserir texto no modelo: {e}")
        return None

if __name__ == '__main__':
    app.run(debug=True)
